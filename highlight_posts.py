#!/usr/bin/env python3

from __future__ import annotations

import argparse
import copy
import sys
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZIP_DEFLATED

try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn
except ImportError as exc:  # pragma: no cover - depends on local environment
    raise SystemExit(
        "python-docx is required. Try running: "
        "/home/weiying/python/word/.venv/bin/python highlight_posts.py"
    ) from exc


@dataclass(frozen=True)
class Segment:
    text: str
    highlight: bool
    source_run_idx: int


def is_current_directory_docx(path: Path) -> bool:
    return path.is_file() and path.suffix.lower() == ".docx" and not path.name.startswith("~$")


def default_targets() -> list[Path]:
    return sorted(
        (path for path in Path.cwd().iterdir() if is_current_directory_docx(path)),
        key=lambda path: path.name.lower(),
    )


def resolve_targets(paths: list[Path]) -> list[Path]:
    if not paths:
        return default_targets()

    targets: list[Path] = []
    for path in paths:
        if path.is_dir():
            targets.extend(
                sorted(
                    (child for child in path.iterdir() if is_current_directory_docx(child)),
                    key=lambda child: child.name.lower(),
                )
            )
        else:
            targets.append(path)
    return targets


def output_path_for(source: Path, output_dir: Path | None, suffix: str) -> Path:
    if output_dir is not None:
        return output_dir / source.name
    if not suffix:
        return source
    return source.with_name(f"{source.stem}{suffix}{source.suffix}")


def run_index_for_offset(run_starts: list[int], run_lengths: list[int], offset: int) -> int:
    if not run_starts:
        return 0
    for idx, start in enumerate(run_starts):
        end = start + run_lengths[idx]
        if start <= offset < end:
            return idx
    return len(run_starts) - 1


def split_star_segments(
    text: str, run_starts: list[int], run_lengths: list[int]
) -> list[Segment] | None:
    segments: list[Segment] = []
    changed = False
    highlight = False
    segment_start = 0

    for idx, char in enumerate(text):
        if char != "*":
            continue

        if idx > segment_start:
            segments.append(
                Segment(
                    text=text[segment_start:idx],
                    highlight=highlight,
                    source_run_idx=run_index_for_offset(
                        run_starts, run_lengths, segment_start
                    ),
                )
            )

        highlight = not highlight
        changed = True
        segment_start = idx + 1

    if not changed:
        return None

    if highlight:
        # Leave paragraphs with unmatched markers untouched.
        return None

    if segment_start < len(text):
        segments.append(
            Segment(
                text=text[segment_start:],
                highlight=False,
                source_run_idx=run_index_for_offset(
                    run_starts, run_lengths, segment_start
                ),
            )
        )

    return [segment for segment in segments if segment.text]


def paragraph_has_hyperlink(paragraph) -> bool:
    return paragraph._p.find(qn("w:hyperlink")) is not None


def clear_direct_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def copy_run_properties(source_run, target_run) -> None:
    source_rpr = source_run._r.rPr
    if source_rpr is None:
        return
    target_rpr = target_run._r.rPr
    if target_rpr is not None:
        target_run._r.remove(target_rpr)
    target_run._r.insert(0, copy.deepcopy(source_rpr))


def highlight_marked_paragraph(paragraph) -> bool:
    if paragraph_has_hyperlink(paragraph):
        return False

    runs = list(paragraph.runs)
    if not runs:
        return False

    run_texts = [run.text for run in runs]
    text = "".join(run_texts)
    run_starts: list[int] = []
    run_lengths: list[int] = []
    offset = 0
    for run_text in run_texts:
        run_starts.append(offset)
        run_lengths.append(len(run_text))
        offset += len(run_text)

    segments = split_star_segments(text, run_starts, run_lengths)
    if not segments:
        return False

    clear_direct_runs(paragraph)
    for segment in segments:
        source_run = runs[min(segment.source_run_idx, len(runs) - 1)]
        new_run = paragraph.add_run(segment.text)
        copy_run_properties(source_run, new_run)
        if segment.highlight:
            new_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    return True


def validate_docx_xml(path: Path) -> None:
    with zipfile.ZipFile(path) as zf:
        ET.fromstring(zf.read("word/document.xml"))


def write_highlighted_package(
    *,
    input_path: Path,
    temp_output_path: Path,
    output_path: Path,
) -> None:
    final_temp_path = output_path.with_suffix(output_path.suffix + ".finaltmp")
    with (
        zipfile.ZipFile(input_path, "r") as zin,
        zipfile.ZipFile(temp_output_path, "r") as ztemp,
        zipfile.ZipFile(final_temp_path, "w", compression=ZIP_DEFLATED) as zout,
    ):
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename == "word/document.xml":
                data = ztemp.read(info.filename)
                ET.fromstring(data)
            zout.writestr(info, data)
    final_temp_path.replace(output_path)


def highlight_docx(source: Path, destination: Path) -> tuple[int, int]:
    doc = Document(str(source))
    changed_paragraphs = 0
    skipped_hyperlink_paragraphs = 0

    for paragraph in doc.paragraphs:
        if "*" not in paragraph.text:
            continue
        if paragraph_has_hyperlink(paragraph):
            skipped_hyperlink_paragraphs += 1
            continue
        if highlight_marked_paragraph(paragraph):
            changed_paragraphs += 1

    if changed_paragraphs:
        destination.parent.mkdir(parents=True, exist_ok=True)
        temp_output_path = destination.with_suffix(destination.suffix + ".tmp")
        doc.save(temp_output_path)
        try:
            write_highlighted_package(
                input_path=source,
                temp_output_path=temp_output_path,
                output_path=destination,
            )
        finally:
            if temp_output_path.exists():
                temp_output_path.unlink()
        validate_docx_xml(destination)

    return changed_paragraphs, skipped_hyperlink_paragraphs


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        prog="highlight-posts",
        description=(
            "Turn paired *text* markers in generated post DOCX files into "
            "bright-green Word highlights."
        )
    )
    parser.add_argument(
        "paths",
        nargs="*",
        type=Path,
        help="DOCX files or folders. Default: all DOCX files in the current directory.",
    )
    parser.add_argument(
        "-o",
        "--output-dir",
        type=Path,
        help="Write processed files to this folder instead of editing in place.",
    )
    parser.add_argument(
        "--suffix",
        default="",
        help=(
            "Suffix for side-by-side output, for example '_highlighted'. "
            "Default is empty, so files are edited in place unless --output-dir is used."
        ),
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Report target files without writing changes.",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv or sys.argv[1:])
    targets = resolve_targets(args.paths)

    if not targets:
        print("[warn] no DOCX files found in the current directory", file=sys.stderr)
        return 1

    exit_code = 0
    for source in targets:
        if not source.exists():
            print(f"[error] not found: {source}", file=sys.stderr)
            exit_code = 1
            continue
        if source.suffix.lower() != ".docx":
            print(f"[skip] not a .docx file: {source}", file=sys.stderr)
            continue

        destination = output_path_for(source, args.output_dir, args.suffix)
        if args.dry_run:
            print(f"[target] {source} -> {destination}")
            continue

        changed, skipped = highlight_docx(source, destination)
        if changed:
            print(f"[updated] {destination} ({changed} paragraph(s))")
        else:
            print(f"[unchanged] {source}")
        if skipped:
            print(
                f"[warn] skipped {skipped} hyperlink paragraph(s) in {source}",
                file=sys.stderr,
            )

    return exit_code


if __name__ == "__main__":
    raise SystemExit(main())
