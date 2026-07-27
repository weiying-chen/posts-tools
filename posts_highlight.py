from __future__ import annotations

import copy
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZIP_DEFLATED

try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn
    from docx.text.run import Run
except ImportError as exc:  # pragma: no cover - depends on local environment
    raise SystemExit(
        "python-docx is required. Try running: "
        "/home/weiying/python/word/.venv/bin/python highlight_posts.py"
    ) from exc


@dataclass(frozen=True)
class Segment:
    text: str
    highlight: bool
    source_run_idx: int


def run_index_for_offset(run_starts: list[int], run_lengths: list[int], offset: int) -> int:
    if not run_starts:
        return 0
    for idx, start in enumerate(run_starts):
        end = start + run_lengths[idx]
        if start <= offset < end:
            return idx
    return len(run_starts) - 1


def split_star_segments(
    text: str, run_starts: list[int], run_lengths: list[int]
) -> list[Segment] | None:
    segments: list[Segment] = []
    changed = False
    highlight = False
    segment_start = 0

    for idx, char in enumerate(text):
        if char != "*":
            continue

        if idx > segment_start:
            segments.append(
                Segment(
                    text=text[segment_start:idx],
                    highlight=highlight,
                    source_run_idx=run_index_for_offset(
                        run_starts, run_lengths, segment_start
                    ),
                )
            )

        highlight = not highlight
        changed = True
        segment_start = idx + 1

    if not changed:
        return None

    if highlight:
        return None

    if segment_start < len(text):
        segments.append(
            Segment(
                text=text[segment_start:],
                highlight=False,
                source_run_idx=run_index_for_offset(
                    run_starts, run_lengths, segment_start
                ),
            )
        )

    return [segment for segment in segments if segment.text]


def split_star_segments_stateful(
    text: str,
    run_starts: list[int],
    run_lengths: list[int],
    highlight: bool,
) -> tuple[list[Segment], bool]:
    segments: list[Segment] = []
    segment_start = 0

    for idx, char in enumerate(text):
        if char != "*":
            continue

        if idx > segment_start:
            segments.append(
                Segment(
                    text=text[segment_start:idx],
                    highlight=highlight,
                    source_run_idx=run_index_for_offset(
                        run_starts, run_lengths, segment_start
                    ),
                )
            )

        highlight = not highlight
        segment_start = idx + 1

    if segment_start < len(text):
        segments.append(
            Segment(
                text=text[segment_start:],
                highlight=highlight,
                source_run_idx=run_index_for_offset(
                    run_starts, run_lengths, segment_start
                ),
            )
        )

    return [segment for segment in segments if segment.text], highlight


def paragraph_has_hyperlink(paragraph) -> bool:
    return paragraph._p.find(qn("w:hyperlink")) is not None


def clear_direct_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def copy_run_properties(source_run, target_run) -> None:
    source_rpr = source_run._r.rPr
    if source_rpr is None:
        return
    target_rpr = target_run._r.rPr
    if target_rpr is not None:
        target_run._r.remove(target_rpr)
    target_run._r.insert(0, copy.deepcopy(source_rpr))


def replace_run_block(paragraph, runs, segments: list[Segment]) -> None:
    first_run_element = runs[0]._element
    insert_at = list(paragraph._p).index(first_run_element)
    for run in runs:
        paragraph._p.remove(run._element)
    for segment in segments:
        source_run = runs[min(segment.source_run_idx, len(runs) - 1)]
        new_run = paragraph.add_run(segment.text)
        copy_run_properties(source_run, new_run)
        if segment.highlight:
            new_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        new_element = new_run._element
        paragraph._p.remove(new_element)
        paragraph._p.insert(insert_at, new_element)
        insert_at += 1


def highlight_run_block(paragraph, runs) -> bool:
    if not runs:
        return False

    run_texts = [run.text for run in runs]
    text = "".join(run_texts)
    run_starts: list[int] = []
    run_lengths: list[int] = []
    offset = 0
    for run_text in run_texts:
        run_starts.append(offset)
        run_lengths.append(len(run_text))
        offset += len(run_text)

    segments = split_star_segments(text, run_starts, run_lengths)
    if not segments:
        return False

    replace_run_block(paragraph, runs, segments)
    return True


def direct_run_blocks(paragraph) -> list[list[Run]]:
    blocks: list[list[Run]] = []
    current: list[Run] = []
    for child in paragraph._p:
        if child.tag == qn("w:r"):
            current.append(Run(child, paragraph))
            continue
        if current:
            blocks.append(current)
            current = []
    if current:
        blocks.append(current)
    return blocks


def highlight_marked_paragraph(paragraph) -> bool:
    changed = False
    for runs in reversed(direct_run_blocks(paragraph)):
        if highlight_run_block(paragraph, runs):
            changed = True
    return changed


def highlight_paragraphs(paragraphs) -> int:
    if sum(paragraph.text.count("*") for paragraph in paragraphs) % 2:
        return sum(highlight_marked_paragraph(paragraph) for paragraph in paragraphs)

    changed_paragraphs = 0
    highlight = False

    for paragraph in paragraphs:
        paragraph_changed = False
        for runs in direct_run_blocks(paragraph):
            run_texts = [run.text for run in runs]
            text = "".join(run_texts)
            if not highlight and "*" not in text:
                continue

            run_starts: list[int] = []
            run_lengths: list[int] = []
            offset = 0
            for run_text in run_texts:
                run_starts.append(offset)
                run_lengths.append(len(run_text))
                offset += len(run_text)

            segments, highlight = split_star_segments_stateful(
                text, run_starts, run_lengths, highlight
            )
            replace_run_block(paragraph, runs, segments)
            paragraph_changed = True

        if paragraph_changed:
            changed_paragraphs += 1

    return changed_paragraphs


def validate_docx_xml(path: Path) -> None:
    with zipfile.ZipFile(path) as zf:
        ET.fromstring(zf.read("word/document.xml"))


def write_highlighted_package(
    *,
    input_path: Path,
    temp_output_path: Path,
    output_path: Path,
) -> None:
    final_temp_path = output_path.with_suffix(output_path.suffix + ".finaltmp")
    with (
        zipfile.ZipFile(input_path, "r") as zin,
        zipfile.ZipFile(temp_output_path, "r") as ztemp,
        zipfile.ZipFile(final_temp_path, "w", compression=ZIP_DEFLATED) as zout,
    ):
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename == "word/document.xml":
                data = ztemp.read(info.filename)
                ET.fromstring(data)
            zout.writestr(info, data)
    final_temp_path.replace(output_path)


def highlight_docx(source: Path, destination: Path) -> tuple[int, int]:
    doc = Document(str(source))
    skipped_hyperlink_paragraphs = 0

    for paragraph in doc.paragraphs:
        if "*" in paragraph.text and paragraph_has_hyperlink(paragraph):
            skipped_hyperlink_paragraphs += 1

    changed_paragraphs = highlight_paragraphs(doc.paragraphs)

    if changed_paragraphs:
        destination.parent.mkdir(parents=True, exist_ok=True)
        temp_output_path = destination.with_suffix(destination.suffix + ".tmp")
        doc.save(temp_output_path)
        try:
            write_highlighted_package(
                input_path=source,
                temp_output_path=temp_output_path,
                output_path=destination,
            )
        finally:
            if temp_output_path.exists():
                temp_output_path.unlink()
        validate_docx_xml(destination)

    return changed_paragraphs, skipped_hyperlink_paragraphs
